# ============================================================
# AI Medical Consultant - Knowledge Base Router
# ============================================================
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.database import KnowledgeDocument, User, get_db
from app.models.schemas import KnowledgeCreate, KnowledgeSearch, KnowledgeResponse
from app.services.rag_service import rag_service
from app.middleware.auth import get_current_user
from app.config import settings

router = APIRouter(prefix="/api/v1/knowledge", tags=["知识库"])


@router.post("/documents")
async def add_knowledge(
    data: KnowledgeCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """添加医学知识文档"""
    # 保存到数据库
    doc = KnowledgeDocument(
        title=data.title,
        category=data.category,
        content=data.content,
        source=data.source,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    # 添加到RAG向量索引
    await rag_service.add_documents(
        texts=[data.content],
        metadatas=[{
            "title": data.title,
            "category": data.category,
            "source": data.source,
        }],
    )

    return {"id": doc.id, "message": "知识文档已添加", "total_docs": rag_service.document_count}


@router.post("/search", response_model=List[KnowledgeResponse])
async def search_knowledge(
    search: KnowledgeSearch,
    current_user: User = Depends(get_current_user),
):
    """搜索医学知识"""
    results = await rag_service.search(
        query=search.query,
        top_k=search.top_k,
        category=search.category,
    )

    return [
        KnowledgeResponse(
            id=r.get("index", i),
            title=r.get("title", ""),
            category=r.get("category", ""),
            content=r.get("content", ""),
            source=r.get("source", ""),
            score=r.get("rrf_score", r.get("score", 0)),
        )
        for i, r in enumerate(results)
    ]


@router.get("/documents")
async def list_documents(
    category: str = None,
    db: AsyncSession = Depends(get_db),
):
    """列出知识文档"""
    query = select(KnowledgeDocument)
    if category:
        query = query.where(KnowledgeDocument.category == category)
    query = query.order_by(KnowledgeDocument.created_at.desc()).limit(100)

    result = await db.execute(query)
    return result.scalars().all()


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    category: str = "general",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """上传医学文档（支持PDF、Word、Markdown、TXT）"""
    import os
    from langchain.text_splitter import RecursiveCharacterTextSplitter

    # 读取文件内容
    content = await file.read()

    # 根据文件类型提取文本
    filename = file.filename or "unknown"
    if filename.endswith(".pdf"):
        from io import BytesIO
        from pypdf2 import PdfReader
        reader = PdfReader(BytesIO(content))
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
    elif filename.endswith(".docx"):
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(content))
        text = "\n".join([p.text for p in doc.paragraphs])
    elif filename.endswith(".md") or filename.endswith(".txt"):
        text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(400, f"不支持的文件格式: {filename}")

    # 文本分块
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", "，", " ", ""],
    )
    chunks = splitter.split_text(text)

    # 保存到数据库和向量索引
    for i, chunk in enumerate(chunks):
        doc = KnowledgeDocument(
            title=f"{filename} (第{i+1}块)",
            category=category,
            content=chunk,
            source=filename,
            chunk_index=i,
        )
        db.add(doc)

    await db.commit()

    # 批量添加到RAG
    await rag_service.add_documents(
        texts=chunks,
        metadatas=[{"title": filename, "category": category, "source": filename} for _ in chunks],
    )

    return {
        "message": f"文档 {filename} 已处理",
        "chunks": len(chunks),
        "total_docs": rag_service.document_count,
    }


@router.get("/stats")
async def knowledge_stats(
    db: AsyncSession = Depends(get_db),
):
    """获取知识库统计信息"""
    from sqlalchemy import func
    result = await db.execute(select(func.count()).select_from(KnowledgeDocument))
    db_count = result.scalar() or 0
    return {
        "total_documents": db_count,
        "index_initialized": rag_service.index is not None,
        "vector_count": rag_service.index.ntotal if rag_service.index else 0,
    }
